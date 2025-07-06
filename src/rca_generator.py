import asyncio
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import json
from src.config import config
from src.mcp_client import mcp_client

from docx import Document
import re

logger = logging.getLogger(__name__)

def extract_template_prompts(template_path):
    """
    Parse the Word template and extract all sections, including headers and their prompts.
    Each section is a dict: {'header': str, 'prompt': str, 'header_idx': int, 'prompt_idx': int}
    Prompts are any text between < > under a header. If no prompt, still create a section with prompt="".
    """
    doc = Document(template_path)
    results = []
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        header = paragraphs[i].text.strip()
        if header:
            # Look for the next paragraph(s) until next header or end
            j = i + 1
            found_prompt = False
            while j < len(paragraphs):
                para_text = paragraphs[j].text.strip()
                prompt_match = re.search(r"<(.+?)>", para_text)
                if prompt_match:
                    prompt = prompt_match.group(1).strip()
                    results.append({
                        "header": header,
                        "prompt": prompt,
                        "header_idx": i,
                        "prompt_idx": j
                    })
                    found_prompt = True
                    break
                # If we hit another header (non-empty, not a prompt), stop
                if para_text and not re.search(r"<(.+?)>", para_text):
                    break
                j += 1
            # If no prompt found, still add the header as a section with empty prompt
            if not found_prompt:
                results.append({
                    "header": header,
                    "prompt": "",
                    "header_idx": i,
                    "prompt_idx": None
                })
        i += 1
    return results

class RCAGenerator:
    def __init__(self):
        self.config = config.llm_config
        self.template_path = Path("data/rca_template.docx")
        self.output_dir = Path(config.app_config['output_directory'])
        self._template_prompts = None
        self._netapp_context = None

    def get_netapp_context(self):
        """Lazily load and cache NetApp context from src/prompts/context_netapp if available."""
        if self._netapp_context is None:
            context_path = Path("src/prompts/context_netapp")
            if context_path.exists():
                with open(context_path, "r", encoding="utf-8") as f:
                    self._netapp_context = f.read().strip()
            else:
                self._netapp_context = ""
        return self._netapp_context

    def get_template_prompts(self):
        """Lazily load and cache prompts from the template docx."""
        if self._template_prompts is None:
            if self.template_path.exists():
                self._template_prompts = extract_template_prompts(self.template_path)
            else:
                self._template_prompts = []
        return self._template_prompts
        
    async def generate_rca_analysis(self, 
                                   files: List[str], 
                                   urls: List[str], 
                                   jira_tickets: List[str],
                                   issue_description: str) -> Dict[str, Any]:
        """Generate RCA analysis based on provided inputs"""
        try:
            logger.info("Starting RCA analysis generation")
            
            # Collect all source data
            source_data = await self._collect_source_data(files, urls, jira_tickets)
            
            # Generate analysis using LLM
            analysis = await self._generate_analysis(source_data, issue_description)

            # --- Inject key fields for template mapping ---
            # 1. Case: look for a support case number (SAP case) in files or filenames
            case_number = None
            for file_path in files:
                # Try to match a long number (e.g., 201012345678)
                import re
                match = re.search(r'\b(20\d{10,})\b', file_path)
                if match:
                    case_number = match.group(1)
                    break
            # 2. CPE: first Jira ticket that matches CPE-xxxx
            cpe_number = None
            for ticket in jira_tickets:
                if ticket.upper().startswith("CPE-"):
                    cpe_number = ticket
                    break
            # 3. Defect: first linked issue that matches CONTAP-xxxxxx or ELEM-xxxxxx
            defect_number = None
            for linked_list in source_data.get("jira_linked_issues", {}).values():
                for issue in linked_list:
                    key = issue.get("key", "")
                    if key.startswith("CONTAP-") or key.startswith("ELEM-"):
                        defect_number = key
                        break
                if defect_number:
                    break
            # Add to analysis if not already present
            if "case" not in analysis and case_number:
                analysis["case"] = case_number
            if "cpe" not in analysis and cpe_number:
                analysis["cpe"] = cpe_number
            if "defect" not in analysis and defect_number:
                analysis["defect"] = defect_number
            # --- End inject ---

            # Add a "sources_used" section to the analysis
            sources_used = []
            if files:
                sources_used.extend([f"File: {Path(f).name}" for f in files])
            if urls:
                sources_used.extend([f"URL: {u}" for u in urls])
            if jira_tickets:
                sources_used.extend([f"Jira Ticket: {jt}" for jt in jira_tickets])
            analysis["sources_used"] = sources_used

            # Create RCA document
            document_path = await self._create_rca_document(analysis)
            
            result = {
                'analysis': analysis,
                'document_path': str(document_path),
                'timestamp': datetime.now().isoformat(),
                'source_data_summary': {
                    'files_processed': len(files),
                    'urls_processed': len(urls),
                    'jira_tickets_referenced': len(jira_tickets)
                }
            }
            
            logger.info("RCA analysis generation completed")
            return result
            
        except Exception as e:
            logger.error(f"Failed to generate RCA analysis: {e}")
            raise
    
    async def _collect_source_data(self, files: List[str], urls: List[str], jira_tickets: List[str]) -> Dict[str, Any]:
        """Collect data from all sources, and for Jira tickets, also fetch linked issues."""
        source_data = {
            'files': {},
            'urls': {},
            'jira_tickets': {},
            'jira_linked_issues': {},
            'summary': ''
        }
        
        # Process files
        for file_path in files:
            try:
                if file_path.lower().endswith('.pdf'):
                    content = await mcp_client.process_pdf(file_path)
                else:
                    content = await mcp_client.read_file(file_path)
                
                source_data['files'][file_path] = {
                    'content': content,
                    'size': len(content),
                    'type': Path(file_path).suffix
                }
                logger.info(f"Processed file: {file_path}")
                
            except Exception as e:
                logger.error(f"Failed to process file {file_path}: {e}")
                source_data['files'][file_path] = {'error': str(e)}
        
        # Process URLs
        for url in urls:
            try:
                content = await mcp_client.scrape_web_content(url)
                source_data['urls'][url] = {
                    'content': content,
                    'size': len(content),
                    'scraped_at': datetime.now().isoformat()
                }
                logger.info(f"Scraped URL: {url}")
                
            except Exception as e:
                logger.error(f"Failed to scrape URL {url}: {e}")
                source_data['urls'][url] = {'error': str(e)}
        
        # Process Jira tickets and collect linked issues
        for ticket_id in jira_tickets:
            try:
                # Search for specific ticket
                jql = f"key = {ticket_id}"
                tickets = await mcp_client.search_jira_tickets(jql, max_results=1)
                
                if tickets:
                    ticket = tickets[0]
                    source_data['jira_tickets'][ticket_id] = ticket
                    logger.info(f"Retrieved Jira ticket: {ticket_id}")

                    # Look for linked issues in the ticket fields
                    linked_issues = []
                    fields = ticket.get('fields', {})
                    issuelinks = fields.get('issuelinks', [])
                    for link in issuelinks:
                        # Outward issue
                        if 'outwardIssue' in link:
                            linked = link['outwardIssue']
                            linked_issues.append({
                                'key': linked.get('key'),
                                'summary': linked.get('fields', {}).get('summary', ''),
                                'type': linked.get('fields', {}).get('issuetype', {}).get('name', ''),
                                'direction': 'outward',
                                'link_type': link.get('type', {}).get('name', '')
                            })
                        # Inward issue
                        if 'inwardIssue' in link:
                            linked = link['inwardIssue']
                            linked_issues.append({
                                'key': linked.get('key'),
                                'summary': linked.get('fields', {}).get('summary', ''),
                                'type': linked.get('fields', {}).get('issuetype', {}).get('name', ''),
                                'direction': 'inward',
                                'link_type': link.get('type', {}).get('name', '')
                            })
                    if linked_issues:
                        source_data['jira_linked_issues'][ticket_id] = linked_issues
                        logger.info(f"Found {len(linked_issues)} linked issues for {ticket_id}")
                else:
                    source_data['jira_tickets'][ticket_id] = {'error': 'Ticket not found'}
                    
            except Exception as e:
                logger.error(f"Failed to retrieve Jira ticket {ticket_id}: {e}")
                source_data['jira_tickets'][ticket_id] = {'error': str(e)}
        
        return source_data
    
    async def _generate_analysis(self, source_data: Dict[str, Any], issue_description: str) -> Dict[str, Any]:
        """Generate RCA analysis using LLM, using the formal prompt from src/prompts/formal_rca_prompt"""
        try:
            # Prepare context for LLM
            context = self._prepare_llm_context(source_data, issue_description)

            # Inject NetApp context if available
            netapp_context = self.get_netapp_context()
            if netapp_context:
                context = f"NETAPP CONTEXT:\n{netapp_context}\n\n{context}"

            # Load the formal RCA prompt from file
            prompt_path = Path("src/prompts/formal_rca_prompt")
            if prompt_path.exists():
                with open(prompt_path, "r", encoding="utf-8") as f:
                    base_prompt = f.read().strip()
                full_prompt = f"{base_prompt}\n\n{context}"
            else:
                # Fallback to default prompt
                full_prompt = (
                    f"Based on the provided context, generate a comprehensive Root Cause Analysis (RCA) report. "
                    f"Structure your response as a JSON object with the following fields:\n"
                    "{\n"
                    "  \"executive_summary\": \"Brief summary of the issue and findings\",\n"
                    "  \"problem_statement\": \"Clear statement of the problem\",\n"
                    "  \"timeline\": \"Chronological sequence of events\",\n"
                    "  \"root_cause\": \"Primary root cause identified\",\n"
                    "  \"contributing_factors\": [\"List of contributing factors\"],\n"
                    "  \"impact_assessment\": \"Assessment of impact\",\n"
                    "  \"corrective_actions\": [\"List of immediate corrective actions\"],\n"
                    "  \"preventive_measures\": [\"List of preventive measures for the future\"],\n"
                    "  \"recommendations\": [\"List of recommendations\"],\n"
                    "  \"escalation_needed\": \"true/false - whether escalation is needed\",\n"
                    "  \"defect_tickets_needed\": \"true/false - whether defect tickets should be created\",\n"
                    "  \"severity\": \"Critical/High/Medium/Low\",\n"
                    "  \"priority\": \"P1/P2/P3/P4\"\n"
                    "}\n"
                    f"\nContext:\n{context}"
                )

            # Generate analysis using configured LLM
            if self.config['default_llm'] == 'openai':
                try:
                    analysis = await self._generate_with_openai(full_prompt)
                except Exception as e:
                    logger.warning(f"OpenAI failed: {e}. Trying fallback...")
                    analysis = await self._try_fallback_llms(['anthropic', 'openrouter'], full_prompt)
            elif self.config['default_llm'] == 'anthropic':
                try:
                    analysis = await self._generate_with_anthropic(full_prompt)
                except Exception as e:
                    logger.warning(f"Anthropic failed: {e}. Trying fallback...")
                    analysis = await self._try_fallback_llms(['openrouter', 'openai'], full_prompt)
            elif self.config['default_llm'] == 'openrouter':
                try:
                    analysis = await self._generate_with_openrouter(full_prompt)
                except Exception as e:
                    logger.warning(f"OpenRouter failed: {e}. Trying fallback...")
                    analysis = await self._try_fallback_llms(['anthropic', 'openai'], full_prompt)
            elif self.config['default_llm'] == 'llmproxy':
                try:
                    analysis = await self._generate_with_llmproxy(full_prompt)
                except Exception as e:
                    logger.warning(f"LLM Proxy failed: {e}. Trying fallback...")
                    analysis = await self._try_fallback_llms(['openai', 'anthropic', 'openrouter'], full_prompt)
            else:
                raise ValueError(f"Unsupported LLM: {self.config['default_llm']}")

            return analysis

        except Exception as e:
            logger.error(f"Failed to generate analysis: {e}")
            raise
    
    def _prepare_llm_context(self, source_data: Dict[str, Any], issue_description: str) -> str:
        """Prepare context string for LLM"""
        context_parts = [
            f"ISSUE DESCRIPTION:\n{issue_description}\n",
            "\n" + "="*50 + "\n",
            "SOURCE DATA ANALYSIS:\n"
        ]
        
        # Add file contents (increase content length limit)
        if source_data['files']:
            context_parts.append("FILES ANALYZED:")
            for file_path, file_data in source_data['files'].items():
                if 'content' in file_data:
                    context_parts.append(f"\n--- {file_path} ---")
                    context_parts.append(file_data['content'][:20000])  # Increased content length
                    if len(file_data['content']) > 20000:
                        context_parts.append("... (content truncated)")
                else:
                    context_parts.append(f"\n--- {file_path} (ERROR) ---")
                    context_parts.append(file_data.get('error', 'Unknown error'))
        
        # Add URL contents (increase content length limit)
        if source_data['urls']:
            context_parts.append("\n\nWEB CONTENT ANALYZED:")
            for url, url_data in source_data['urls'].items():
                if 'content' in url_data:
                    context_parts.append(f"\n--- {url} ---")
                    context_parts.append(url_data['content'][:10000])  # Increased content length
                    if len(url_data['content']) > 10000:
                        context_parts.append("... (content truncated)")
                else:
                    context_parts.append(f"\n--- {url} (ERROR) ---")
                    context_parts.append(url_data.get('error', 'Unknown error'))
        
        # Add Jira ticket data
        if source_data['jira_tickets']:
            context_parts.append("\n\nJIRA TICKETS REFERENCED:")
            for ticket_id, ticket_data in source_data['jira_tickets'].items():
                if 'key' in ticket_data:
                    context_parts.append(f"\n--- {ticket_id} ---")
                    context_parts.append(f"Summary: {ticket_data['summary']}")
                    context_parts.append(f"Status: {ticket_data['status']}")
                    context_parts.append(f"Priority: {ticket_data['priority']}")
                    context_parts.append(f"Description: {ticket_data['description']}")
                else:
                    context_parts.append(f"\n--- {ticket_id} (ERROR) ---")
                    context_parts.append(ticket_data.get('error', 'Unknown error'))
        
        return "\n".join(context_parts)
    
    async def _generate_with_openai(self, context: str) -> Dict[str, Any]:
        """Generate analysis using OpenAI"""
        try:
            import openai
            
            client = openai.AsyncOpenAI(
                api_key=self.config['openai_api_key'],
                base_url=self.config['openai_base_url']
            )
            
            prompt = f"""
            Based on the provided context, generate a comprehensive Root Cause Analysis (RCA) report. 
            Structure your response as a JSON object with the following fields:
            
            {{
                "executive_summary": "Brief summary of the issue and findings",
                "problem_statement": "Clear statement of the problem",
                "timeline": "Chronological sequence of events",
                "root_cause": "Primary root cause identified",
                "contributing_factors": ["List of contributing factors"],
                "impact_assessment": "Assessment of impact",
                "corrective_actions": ["List of immediate corrective actions"],
                "preventive_measures": ["List of preventive measures for the future"],
                "recommendations": ["List of recommendations"],
                "escalation_needed": "true/false - whether escalation is needed",
                "defect_tickets_needed": "true/false - whether defect tickets should be created",
                "severity": "Critical/High/Medium/Low",
                "priority": "P1/P2/P3/P4"
            }}
            
            Context:
            {context}
            """
            
            response = await client.chat.completions.create(
                model=self.config['openai_model'],
                messages=[
                    {"role": "system", "content": "You are an expert technical analyst specializing in root cause analysis."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            # Parse JSON response
            analysis_text = response.choices[0].message.content
            logger.debug(f"OpenAI response content: {analysis_text}")
            
            if not analysis_text or analysis_text.strip() == "":
                logger.error("OpenAI returned empty response")
                raise ValueError("Empty response from OpenAI")
            
            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON response: {e}")
                logger.error(f"Raw response: {analysis_text}")
                
                # Try to extract JSON from response if it's wrapped in text
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    try:
                        analysis = json.loads(json_match.group())
                        logger.info("Successfully extracted JSON from response")
                    except json.JSONDecodeError:
                        # Fallback: create a basic analysis structure
                        analysis = self._create_fallback_analysis(analysis_text)
                else:
                    # Fallback: create a basic analysis structure
                    analysis = self._create_fallback_analysis(analysis_text)
            
            return analysis
            
        except Exception as e:
            logger.error(f"Failed to generate analysis with OpenAI: {e}")
            raise
    
    async def _generate_with_anthropic(self, context: str) -> Dict[str, Any]:
        """Generate analysis using Anthropic"""
        try:
            import anthropic
            
            client = anthropic.AsyncAnthropic(
                api_key=self.config['anthropic_api_key']
            )
            
            prompt = f"""
            Based on the provided context, generate a comprehensive Root Cause Analysis (RCA) report. 
            Structure your response as a JSON object with the following fields:
            
            {{
                "executive_summary": "Brief summary of the issue and findings",
                "problem_statement": "Clear statement of the problem",
                "timeline": "Chronological sequence of events",
                "root_cause": "Primary root cause identified",
                "contributing_factors": ["List of contributing factors"],
                "impact_assessment": "Assessment of impact",
                "corrective_actions": ["List of immediate corrective actions"],
                "preventive_measures": ["List of preventive measures for the future"],
                "recommendations": ["List of recommendations"],
                "escalation_needed": "true/false - whether escalation is needed",
                "defect_tickets_needed": "true/false - whether defect tickets should be created",
                "severity": "Critical/High/Medium/Low",
                "priority": "P1/P2/P3/P4"
            }}
            
            Context:
            {context}
            """
            
            response = await client.messages.create(
                model=self.config['anthropic_model'],
                max_tokens=4000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Parse JSON response
            analysis_text = response.content[0].text
            logger.debug(f"Anthropic response content: {analysis_text}")
            
            if not analysis_text or analysis_text.strip() == "":
                logger.error("Anthropic returned empty response")
                raise ValueError("Empty response from Anthropic")
            
            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON response: {e}")
                logger.error(f"Raw response: {analysis_text}")
                
                # Try to extract JSON from response if it's wrapped in text
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    try:
                        analysis = json.loads(json_match.group())
                        logger.info("Successfully extracted JSON from response")
                    except json.JSONDecodeError:
                        # Fallback: create a basic analysis structure
                        analysis = self._create_fallback_analysis(analysis_text)
                else:
                    # Fallback: create a basic analysis structure
                    analysis = self._create_fallback_analysis(analysis_text)
            
            return analysis
            
        except Exception as e:
            logger.error(f"Failed to generate analysis with Anthropic: {e}")
            raise
    
    async def _generate_with_openrouter(self, context: str) -> Dict[str, Any]:
        """Generate analysis using OpenRouter"""
        try:
            import openai
            
            client = openai.AsyncOpenAI(
                api_key=self.config['openrouter_api_key'],
                base_url=self.config['openrouter_base_url']
            )
            
            prompt = f"""
            Based on the provided context, generate a comprehensive Root Cause Analysis (RCA) report. 
            Structure your response as a JSON object with the following fields:
            
            {{
                "executive_summary": "Brief summary of the issue and findings",
                "problem_statement": "Clear statement of the problem",
                "timeline": "Chronological sequence of events",
                "root_cause": "Primary root cause identified",
                "contributing_factors": ["List of contributing factors"],
                "impact_assessment": "Assessment of impact",
                "corrective_actions": ["List of immediate corrective actions"],
                "preventive_measures": ["List of preventive measures for the future"],
                "recommendations": ["List of recommendations"],
                "escalation_needed": "true/false - whether escalation is needed",
                "defect_tickets_needed": "true/false - whether defect tickets should be created",
                "severity": "Critical/High/Medium/Low",
                "priority": "P1/P2/P3/P4"
            }}
            
            Context:
            {context}
            """
            
            response = await client.chat.completions.create(
                model=self.config['openrouter_model'],
                messages=[
                    {"role": "system", "content": "You are an expert technical analyst specializing in root cause analysis."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000,
                extra_headers={
                    "HTTP-Referer": "https://github.com/solarisjon/creatercav4",
                    "X-Title": "MCP RCA Tool"
                }
            )
            
            # Parse JSON response
            analysis_text = response.choices[0].message.content
            logger.debug(f"OpenRouter response content: {analysis_text}")
            
            if not analysis_text or analysis_text.strip() == "":
                logger.error("OpenRouter returned empty response")
                raise ValueError("Empty response from OpenRouter")
            
            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON response: {e}")
                logger.error(f"Raw response: {analysis_text}")
                
                # Try to extract JSON from response if it's wrapped in text
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    try:
                        analysis = json.loads(json_match.group())
                        logger.info("Successfully extracted JSON from response")
                    except json.JSONDecodeError:
                        # Fallback: create a basic analysis structure
                        analysis = self._create_fallback_analysis(analysis_text)
                else:
                    # Fallback: create a basic analysis structure
                    analysis = self._create_fallback_analysis(analysis_text)
            
            return analysis
            
        except Exception as e:
            logger.error(f"Failed to generate analysis with OpenRouter: {e}")
            raise
    
    async def _generate_with_llmproxy(self, context: str) -> Dict[str, Any]:
        """Generate analysis using LLM Proxy (OpenAI-compatible)"""
        try:
            import openai
            
            client = openai.AsyncOpenAI(
                api_key=self.config['llmproxy_api_key'],
                base_url=self.config['llmproxy_base_url']
            )
            
            prompt = f"""
            Based on the provided context, generate a comprehensive Root Cause Analysis (RCA) report. 
            Structure your response as a JSON object with the following fields:
            
            {{
                "executive_summary": "Brief summary of the issue and findings",
                "problem_statement": "Clear statement of the problem",
                "timeline": "Chronological sequence of events",
                "root_cause": "Primary root cause identified",
                "contributing_factors": ["List of contributing factors"],
                "impact_assessment": "Assessment of impact",
                "corrective_actions": ["List of immediate corrective actions"],
                "preventive_measures": ["List of preventive measures for the future"],
                "recommendations": ["List of recommendations"],
                "escalation_needed": "true/false - whether escalation is needed",
                "defect_tickets_needed": "true/false - whether defect tickets should be created",
                "severity": "Critical/High/Medium/Low",
                "priority": "P1/P2/P3/P4"
            }}
            
            Context:
            {context}
            """
            
            response = await client.chat.completions.create(
                model=self.config['llmproxy_model'],
                messages=[
                    {"role": "system", "content": "You are an expert technical analyst specializing in root cause analysis."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            # Parse JSON response
            analysis_text = response.choices[0].message.content
            logger.debug(f"LLM Proxy response content: {analysis_text}")
            
            if not analysis_text or analysis_text.strip() == "":
                logger.error("LLM Proxy returned empty response")
                raise ValueError("Empty response from LLM Proxy")
            
            try:
                analysis = json.loads(analysis_text)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON response: {e}")
                logger.error(f"Raw response: {analysis_text}")
                
                # Try to extract JSON from response if it's wrapped in text
                import re
                json_match = re.search(r'\{.*\}', analysis_text, re.DOTALL)
                if json_match:
                    try:
                        analysis = json.loads(json_match.group())
                        logger.info("Successfully extracted JSON from response")
                    except json.JSONDecodeError:
                        # Fallback: create a basic analysis structure
                        analysis = self._create_fallback_analysis(analysis_text)
                else:
                    # Fallback: create a basic analysis structure
                    analysis = self._create_fallback_analysis(analysis_text)
            
            return analysis
            
        except Exception as e:
            logger.error(f"Failed to generate analysis with LLM Proxy: {e}")
            raise
    
    async def _try_fallback_llms(self, fallback_order: List[str], context: str) -> Dict[str, Any]:
        """Try fallback LLMs in order until one succeeds"""
        for llm_name in fallback_order:
            try:
                logger.info(f"Trying fallback LLM: {llm_name}")
                
                if llm_name == 'openai' and self.config.get('openai_api_key'):
                    return await self._generate_with_openai(context)
                elif llm_name == 'anthropic' and self.config.get('anthropic_api_key'):
                    return await self._generate_with_anthropic(context)
                elif llm_name == 'openrouter' and self.config.get('openrouter_api_key'):
                    return await self._generate_with_openrouter(context)
                elif llm_name == 'llmproxy' and self.config.get('llmproxy_api_key'):
                    return await self._generate_with_llmproxy(context)
                else:
                    logger.warning(f"No API key available for {llm_name}")
                    continue
                    
            except Exception as e:
                logger.warning(f"Fallback LLM {llm_name} failed: {e}")
                continue
        
        # If all fallbacks fail, raise the last exception
        raise Exception("All LLM providers failed to generate analysis")
    
    def _create_fallback_analysis(self, raw_text: str) -> Dict[str, Any]:
        """Create a fallback analysis structure when JSON parsing fails"""
        logger.warning("Using fallback analysis structure due to JSON parsing error")
        
        return {
            "executive_summary": "RCA analysis generated with limited parsing. Please review the raw analysis below.",
            "problem_statement": "Unable to parse structured analysis from LLM response.",
            "timeline": "Timeline information not available in structured format.",
            "root_cause": "Root cause analysis requires manual review of the raw response.",
            "contributing_factors": ["JSON parsing failure", "LLM response format issue"],
            "impact_assessment": "Impact assessment requires manual review.",
            "corrective_actions": ["Review raw LLM response", "Adjust prompt formatting", "Check API configuration"],
            "preventive_measures": ["Improve response parsing", "Add response validation"],
            "recommendations": ["Manual review of raw analysis", "Check LLM API configuration"],
            "escalation_needed": "false",
            "defect_tickets_needed": "true",
            "severity": "Medium",
            "priority": "P3",
            "raw_response": raw_text[:2000]  # Include first 2000 chars of raw response
        }
    
    async def _create_rca_document(self, analysis: Dict[str, Any]) -> Path:
        """Create RCA document from analysis, filling in the Word template.
        Prompts are detected as <...> and replaced with generated text.
        If not enough data, fill with 'Unable to accurately find enough data'.
        """
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = self.output_dir / f"rca_report_{timestamp}.docx"
            self.output_dir.mkdir(parents=True, exist_ok=True)

            template_path = Path("rca_template_doc.docx")
            if not template_path.exists():
                # fallback to JSON if template is missing
                json_file = self.output_dir / f"rca_report_{timestamp}.json"
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(analysis, f, indent=2, ensure_ascii=False)
                logger.warning(f"Template not found, saved as JSON: {json_file}")
                return json_file

            doc = Document(template_path)

            # Extract all <prompt> under each header
            prompts = extract_template_prompts(template_path)

            # Build a mapping from header to prompt and paragraph index
            header_to_prompt_idx = {}
            for section in prompts:
                header = section['header']
                prompt = section['prompt']
                header_idx = section['header_idx']
                prompt_idx = section['prompt_idx']
                header_to_prompt_idx[header] = {
                    "prompt": prompt,
                    "header_idx": header_idx,
                    "prompt_idx": prompt_idx
                }

            # Improved: Use a synonym map and fuzzy matching to map headers to analysis keys
            def find_analysis_key(header):
                norm_header = header.lower().replace(" ", "_")
                synonym_map = {
                    "customer": ["customer", "customer_name", "client", "account"],
                    "cases": ["cases", "case", "case_number", "support_case", "support_cases", "tickets", "jira_tickets"],
                    "synopsis": ["synopsis", "summary", "executive_summary", "overview", "description"],
                    "issue_tracking_number": ["issue_tracking_number", "case", "case_number", "cases", "support_case", "sap_case", "tracking_number"],
                    "cpe": ["cpe", "cpe_number"],
                    "defect": ["defect", "defect_ids", "defects", "related_defects"],
                    "cap_color": ["cap_color", "cap", "color"],
                    "timeline": ["timeline"],
                    "executive_summary": ["executive_summary", "summary", "overview"],
                    "problem_summary": ["problem_summary", "problem_statement", "problem", "issue_summary"],
                    "impact": ["impact", "impact_assessment"],
                    "root_cause": ["root_cause", "cause"],
                    "likelihood_of_occurrence": ["likelihood_of_occurrence", "likelihood", "probability"],
                    "vulnerability": ["vulnerability", "vulnerabilities"],
                    "overall_risk_profile": ["overall_risk_profile", "risk_profile", "risk"],
                    "workaround": ["workaround", "workarounds"],
                    "known_defects_and_resolution": ["known_defects_and_resolution", "known_defects", "defect_resolution"],
                    "new_defects_and_resolution": ["new_defects_and_resolution", "new_defects"],
                    "recommended_changes": ["recommended_changes", "recommendations", "recommended_system_changes"],
                    "prevention_current": ["prevention_current", "prevention", "current_prevention"],
                    "prevention_future": ["prevention_future", "future_prevention"],
                    "monitoring": ["monitoring", "monitor", "monitor_for_prevention"],
                }
                # Try direct match
                for k in analysis.keys():
                    if k.lower().replace(" ", "_") == norm_header:
                        return k
                # Try synonyms
                for syn_header, syn_list in synonym_map.items():
                    if norm_header == syn_header:
                        for syn in syn_list:
                            for k in analysis.keys():
                                if k.lower().replace(" ", "_") == syn:
                                    return k
                # Try partial/fuzzy match
                for k in analysis.keys():
                    if norm_header in k.lower().replace(" ", "_") or k.lower().replace(" ", "_") in norm_header:
                        return k
                return None

            # For each template section, fill the value after the header, always, in order.
            for section in prompts:
                header = section['header']
                prompt_idx = section['prompt_idx']
                header_idx = section['header_idx']
                key = find_analysis_key(header)
                value = analysis.get(key) if key else None
                if value is None or (isinstance(value, str) and not value.strip()):
                    value = "Unable to accurately find enough data"
                elif isinstance(value, list):
                    value = "\n".join(f"- {v}" for v in value)
                elif isinstance(value, bool):
                    value = "Yes" if value else "No"
                else:
                    value = str(value)
                # If prompt is present, replace it
                if prompt_idx is not None:
                    para = doc.paragraphs[prompt_idx]
                    para.text = re.sub(r"<(.+?)>", value, para.text)
                    # If no <...> found, just replace the whole paragraph with value
                    if "<" not in para.text and ">" not in para.text and not re.search(r"<(.+?)>", para.text):
                        para.text = value
                # Always insert value after the header (even if prompt is present or not)
                if header_idx is not None and header_idx + 1 < len(doc.paragraphs):
                    doc.paragraphs[header_idx + 1].text = value

            # Save the filled document
            doc.save(output_file)
            logger.info(f"RCA Word document created: {output_file}")
            return output_file

        except Exception as e:
            logger.error(f"Failed to create RCA document: {e}")
            raise

# Global RCA generator instance
rca_generator = RCAGenerator()
