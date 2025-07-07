"""
Core RCA Analysis Engine
Orchestrates the entire analysis process using the refactored components
"""
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from ..llm.client import UnifiedLLMClient
from .prompt_manager import PromptManager
from .parsers import ResponseParser

logger = logging.getLogger(__name__)

class RCAEngine:
    """Main analysis engine that orchestrates the RCA process"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.llm_client = UnifiedLLMClient(config)
        self.prompt_manager = PromptManager()
        self.response_parser = ResponseParser()
        self.mcp_client = None  # Will be injected
        
        self.output_dir = Path(config.get('output_directory', 'output'))
        self.output_dir.mkdir(exist_ok=True)
    
    def set_mcp_client(self, mcp_client):
        """Inject MCP client dependency"""
        self.mcp_client = mcp_client
    
    async def generate_analysis(self, 
                              files: List[str],
                              urls: List[str], 
                              jira_tickets: List[str],
                              analysis_type: str = "formal_rca_prompt",
                              issue_description: str = "") -> Dict[str, Any]:
        """
        Generate comprehensive RCA analysis
        
        Args:
            files: List of file paths to analyze
            urls: List of URLs to analyze
            jira_tickets: List of Jira ticket IDs
            issue_description: User's description of the issue
            prompt_type: Type of analysis to perform
            
        Returns:
            Analysis result dictionary
        """
        try:
            logger.info(f"Starting {analysis_type} analysis")
            
            # Collect source data
            source_data = await self._collect_source_data(files, urls, jira_tickets)
            
            # Build prompt
            prompt = self.prompt_manager.build_prompt(
                prompt_type=analysis_type,
                context_data=source_data,
                issue_description=issue_description
            )
            
            # Generate analysis using LLM
            raw_response = await self.llm_client.generate_analysis(
                prompt=prompt,
                preferred_provider=self.config.get('default_llm')
            )
            
            # Parse response
            parsed_analysis = self.response_parser.parse_llm_response(raw_response, analysis_type)
            
            # Add metadata
            result = {
                'analysis': parsed_analysis,
                'prompt_file_used': analysis_type,
                'timestamp': datetime.now().isoformat(),
                'sources_count': {
                    'files': len(files),
                    'urls': len(urls), 
                    'jira_tickets': len(jira_tickets)
                }
            }
            
            # Add sources used to analysis
            if 'sources_used' not in parsed_analysis:
                parsed_analysis['sources_used'] = self._build_sources_list(files, urls, jira_tickets)
            
            # For formal RCA, generate document
            if analysis_type == "formal_rca_prompt":
                doc_path = await self._generate_document(result)
                result['document_path'] = str(doc_path)
            
            # Save JSON report
            json_path = await self._save_json_report(result, analysis_type)
            result['json_path'] = str(json_path)
            
            logger.info(f"Analysis completed successfully: {analysis_type}")
            return result
            
        except Exception as e:
            logger.error(f"Analysis generation failed: {e}")
            raise
    
    async def _collect_source_data(self, files: List[str], urls: List[str], 
                                 jira_tickets: List[str]) -> str:
        """Collect and combine data from all sources"""
        if not self.mcp_client:
            raise ValueError("MCP client not initialized")
        
        data_parts = []
        
        # Process files
        if files:
            logger.info(f"Processing {len(files)} files")
            for file_path in files:
                try:
                    content = await self.mcp_client.read_file(file_path)
                    data_parts.append(f"## File: {file_path}\n{content}\n")
                except Exception as e:
                    logger.warning(f"Failed to read file {file_path}: {e}")
        
        # Process URLs
        if urls:
            logger.info(f"Processing {len(urls)} URLs")
            for url in urls:
                try:
                    content = await self.mcp_client.scrape_web_content(url)
                    data_parts.append(f"## URL: {url}\n{content}\n")
                except Exception as e:
                    logger.warning(f"Failed to fetch URL {url}: {e}")
        
        # Process Jira tickets
        if jira_tickets:
            logger.info(f"Processing {len(jira_tickets)} Jira tickets")
            for ticket_id in jira_tickets:
                try:
                    ticket_data = await self.mcp_client.get_jira_ticket(ticket_id)
                    data_parts.append(f"## Jira Ticket: {ticket_id}\n{ticket_data}\n")
                except Exception as e:
                    logger.warning(f"Failed to fetch Jira ticket {ticket_id}: {e}")
        
        if not data_parts:
            return "No source data available for analysis."
        
        return "\n".join(data_parts)
    
    def _build_sources_list(self, files: List[str], urls: List[str], 
                          jira_tickets: List[str]) -> List[str]:
        """Build list of sources used in analysis"""
        sources = []
        
        for file_path in files:
            sources.append(f"File: {Path(file_path).name}")
        
        for url in urls:
            sources.append(f"URL: {url}")
        
        for ticket in jira_tickets:
            sources.append(f"Jira: {ticket}")
        
        return sources
    
    async def _generate_document(self, analysis_result: Dict[str, Any]) -> Path:
        """Generate Word document for formal RCA"""
        from docx import Document
        
        doc = Document()
        analysis = analysis_result['analysis']
        
        # Title
        doc.add_heading('Root Cause Analysis Report', 0)
        doc.add_paragraph(f"Generated: {analysis_result['timestamp']}")
        
        # Sections
        sections = [
            ('Executive Summary', 'executive_summary'),
            ('Problem Statement', 'problem_statement'),
            ('Timeline', 'timeline'),
            ('Root Cause', 'root_cause'),
            ('Contributing Factors', 'contributing_factors'),
            ('Impact Assessment', 'impact_assessment'),
            ('Corrective Actions', 'corrective_actions'),
            ('Preventive Measures', 'preventive_measures'),
            ('Recommendations', 'recommendations'),
        ]
        
        for title, key in sections:
            if key in analysis and analysis[key]:
                doc.add_heading(title, 1)
                content = analysis[key]
                if isinstance(content, list):
                    for item in content:
                        doc.add_paragraph(f"• {item}")
                else:
                    doc.add_paragraph(str(content))
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = self.output_dir / f"rca_report_{timestamp}.docx"
        doc.save(str(doc_path))
        
        logger.info(f"Document saved: {doc_path}")
        return doc_path
    
    async def _save_json_report(self, analysis_result: Dict[str, Any], 
                              prompt_type: str) -> Path:
        """Save JSON report"""
        import json
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{prompt_type}_{timestamp}.json"
        json_path = self.output_dir / filename
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis_result, f, indent=2, ensure_ascii=False)
        
        logger.info(f"JSON report saved: {json_path}")
        return json_path
    
    def get_available_analysis_types(self) -> List[str]:
        """Get list of available analysis types"""
        return self.prompt_manager.get_available_prompts()
    
    def validate_configuration(self) -> Dict[str, bool]:
        """Validate that all required components are properly configured"""
        validation = {
            'llm_providers': len(self.llm_client.get_available_providers()) > 0,
            'prompts_available': len(self.get_available_analysis_types()) > 0,
            'output_directory': self.output_dir.exists() and self.output_dir.is_dir(),
            'mcp_client': self.mcp_client is not None
        }
        
        logger.info(f"Configuration validation: {validation}")
        return validation
